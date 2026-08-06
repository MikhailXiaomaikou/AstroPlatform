#!/usr/bin/env python3
"""Validate the generated Standard Astro v0.2 formal report package (rev 1.1).

Beyond structural checks, this validator enforces the honesty contract of
revision 1.1: the spec-language matrix's perfect system score must always be
labeled as a deterministic-path self-check with the model not in the loop,
the withdrawn per-model "system gain" framing must not reappear, and the
natural-phrasing matrix with its model-participation strata must be present.
"""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[2]
REPORT_ROOT = REPO_ROOT / "docs/research/formal_report_package_2026-08-06"
SCORES = REPORT_ROOT / "evidence/standard_astro_v02_scores_240.csv"
NATURAL_SCORES = REPORT_ROOT / "evidence/standard_astro_v02_natural_scores_240.csv"
POSTFIX_SCORES = REPORT_ROOT / "evidence/standard_astro_v02_natural_postfix_scores_240.csv"

# Withdrawn framings that must never come back in any report body.
FORBIDDEN_EVERYWHERE = (
    "系统增益",
    "system gain",
    "系统辅助样本",
)


def document_text(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def main() -> None:
    docx_files = sorted(REPORT_ROOT.rglob("*.docx"))
    assert len(docx_files) == 10, f"expected 10 DOCX files, found {len(docx_files)}"

    with SCORES.open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    assert len(rows) == 240, f"expected 240 score rows, found {len(rows)}"

    with NATURAL_SCORES.open(encoding="utf-8", newline="") as handle:
        natural_rows = list(csv.DictReader(handle))
    assert len(natural_rows) == 240, (
        f"expected 240 natural score rows, found {len(natural_rows)}"
    )
    assert "llm_calls" in natural_rows[0], "natural scores must record llm_calls"

    with POSTFIX_SCORES.open(encoding="utf-8", newline="") as handle:
        postfix_rows = list(csv.DictReader(handle))
    assert len(postfix_rows) == 240, (
        f"expected 240 post-fix score rows, found {len(postfix_rows)}"
    )
    assert "stratum" in natural_rows[0], "natural scores must record stratum"
    strata = {row["stratum"] for row in natural_rows if row["condition"] == "standard_astro"}
    assert strata <= {"standard_pipeline", "standard_model_in_loop"}, strata

    all_texts: dict[str, str] = {}
    for path in docx_files:
        all_texts[path.name] = document_text(path)
    for name, text in all_texts.items():
        for forbidden in FORBIDDEN_EVERYWHERE:
            assert forbidden not in text, f"{name} contains withdrawn framing {forbidden!r}"

    technical = all_texts["01_Standard_Astro_v0.2_总体技术报告.docx"]
    for expected in (
        "总体技术报告",
        "deterministic_source_check",
        "research_exploration",
        "full_research",
        "general",
        "1440/1440",
        "模型不在回路",
        "不构成模型行为证据",
        "自然措辞矩阵",
        "博士后 12 组匿名 A/B 盲评尚未完成",
    ):
        assert expected in technical, f"technical report missing {expected!r}"

    summary = all_texts["02_Standard_Astro_v0.2_测试结果综述.docx"]
    for expected in (
        "839/1440",
        "58.3%",
        "确定性路径自检",
        "模型不在回路",
        "自然措辞矩阵",
        "模型在回路",
        "llm_calls",
        "B1/B3/B4",
        "Kimi K3",
        "已撤回",
        "误差预算表",
        "95% 置信上界",
        "should-pass",
        "人工判读",
        "子串误报",
        "8.2 缺陷修复与验证复跑",
        "修复后",
        "refusal×15",
    ):
        assert expected in summary, f"evaluation summary missing {expected!r}"

    experiment_docs = sorted((REPORT_ROOT / "03_逐实验报告").glob("*.docx"))
    assert len(experiment_docs) == 8
    for index, path in enumerate(experiment_docs, 1):
        text = all_texts[path.name]
        for expected in (
            f"实验 {index}",
            "科学背景与研究问题",
            "预注册验收标准",
            "测试结果",
            "系统行为审计",
            "自然措辞矩阵对照",
            "模型不在回路",
            "结论范围与局限",
            "English abstract",
        ):
            assert expected in text, f"{path.name} missing {expected!r}"

    required_evidence = {
        "standard_astro_v02_preregistered_tasks.json",
        "standard_astro_v02_scores_240.csv",
        "standard_astro_v02_summary.json",
        "standard_astro_v02_natural_preregistered_tasks.json",
        "standard_astro_v02_natural_scores_240.csv",
        "standard_astro_v02_natural_summary.json",
        "standard_astro_v02_should_pass_corpus.json",
        "standard_astro_v02_should_pass_corpus_postfix.json",
        "standard_astro_v02_natural_postfix_scores_240.csv",
        "standard_astro_v02_natural_postfix_summary.json",
        "strict_blind_test_summary.md",
    }
    actual_evidence = {path.name for path in (REPORT_ROOT / "evidence").iterdir()}
    assert required_evidence <= actual_evidence, required_evidence - actual_evidence

    readme = (REPORT_ROOT / "README.md").read_text(encoding="utf-8")
    for expected in ("修订 1.2", "修订 1.1", "已撤回", "模型不在回路", "模型在回路层", "95% 置信上界", "硬门失守 0"):
        assert expected in readme, f"README missing {expected!r}"

    print(
        "PASS: 10 DOCX reports, 3x240 score rows, honesty relabels present, "
        "withdrawn framings absent, error budget, CI phrasing and post-fix "
        "verification present, 11 evidence artifacts validated."
    )


if __name__ == "__main__":
    main()
